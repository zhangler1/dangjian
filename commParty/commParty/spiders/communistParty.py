import scrapy
from bs4 import BeautifulSoup
from docx import Document
import  re

class CommunistpartySpider(scrapy.Spider):
    name = "communistParty"
    preTitle=""
    def start_requests(self):
        # 定义起始URL
        urls = [
            'https://www.12371.cn/special/xjpznd/'
        ]
        for url in urls:
            yield scrapy.Request(url, callback=self.parseContents)

    def parseContents(self, response):

        soup = BeautifulSoup(response.text, 'lxml')

        hyperlinks = soup.find_all('a')
        for aTag in hyperlinks:
            text = aTag.text
            if "href" in aTag.attrs.keys():
                url = aTag.attrs["href"]
                if re.match(r"（", text):
                    title = self.preTitle + aTag.text
                else:
                    title =aTag.text
                    self.preTitle = text
                if re.search(f"ARTI", url):
                    yield scrapy.Request(url, callback=self.parseArticle, meta={'title': title})




    def parseArticle(self, response):
        soup = BeautifulSoup(response.text, 'lxml')
        fulltext = soup.find('div', {'class': 'dangyuanwang160317_ind01'})
        fileName = response.meta["title"]
        title = fulltext.find('h1', {'class': 'big_title'})
        fulltext = fulltext.find('div', {'class': 'font_area_mid'})
        # 创建一个新的 Word 文档
        doc = Document()

        # 添加标题
        titleStyle = doc.add_heading(title.text, level=1)

        mainText = fulltext.find_all('p')
        for p in mainText:
            if not re.search("延伸阅读", p.text):
                doc.add_paragraph(f'{p.text}')
        # 设置标题的样式（可选）
        titleStyle.bold = True
        titleStyle.italic = False
        titleStyle.underline = False
        fileName = fileName.replace("|", "_").strip()
        fileName = fileName.replace("，", "_")
        fileName = fileName.replace("*", "_")
        fileName = fileName.replace(">", "")
        fileName = fileName.replace("<", "")
        # 保存文档到本地，
        doc.save(f"{fileName}.docx")

class CommunistpartyHistorySpider(scrapy.Spider):
    name = "communistPartyHistory"
    preTitle=""
    def start_requests(self):
        # 定义起始URL
        urls = [
            'https://www.neac.gov.cn/seac/c100907/jiandang90_list.shtml'
        ]
        for url in urls:
            yield scrapy.Request(url, callback=self.parseContents)

    def parseContents(self, response):

        soup = BeautifulSoup(response.text, 'lxml')

        hyperlinks = soup.find_all('a')
        for aTag in hyperlinks:
            text = aTag.text
            if "href" in aTag.attrs.keys():
                url = aTag.attrs["href"]
                if re.match(r"（", text):
                    title = self.preTitle + aTag.text
                else:
                    title =aTag.text
                    self.preTitle = text
                if url:
                    url="https://www.neac.gov.cn"+url
                    yield scrapy.Request(url, callback=self.parseArticle, meta={'title': title})




    def parseArticle(self, response):
        soup = BeautifulSoup(response.text, 'lxml')

        fulltext = soup.find('div', {'class': 'f1-3-1'})
        if not fulltext:
            return
        fileName = response.meta["title"]

        title = fulltext.find('p', {'class': 'p1'})
        fulltext = fulltext.find('div', {'class': 'p3'})
        # 创建一个新的 Word 文档
        doc = Document()

        # 添加标题
        titleStyle = doc.add_heading(title.text, level=1)

        mainText = fulltext.find_all('p')
        for p in mainText:
            if not re.search("延伸阅读", p.text):
                doc.add_paragraph(f'{p.text}')
        # 设置标题的样式（可选）
        titleStyle.bold = True
        titleStyle.italic = False
        titleStyle.underline = False
        fileName = fileName.replace("|", "_").strip()
        fileName = fileName.replace("，", "_")
        fileName = fileName.replace("*", "_")
        fileName = fileName.replace(">", "")
        fileName = fileName.replace("<", "")
        # 保存文档到本地，
        doc.save(f"{fileName}.docx")